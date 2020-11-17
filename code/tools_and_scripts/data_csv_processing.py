import os, sys
# from subprocess import run
from tkinter.filedialog import askopenfilename
import pandas as pd 
import postcode_finder as pf
from copy import copy
from docx import Document
from docx.enum.section import WD_ORIENT
from datetime import datetime

def df_to_word_table(filename, df, text_before=''):
	# Mostly taken from:
	# https://stackoverflow.com/questions/40596518/writing-a-python-pandas-dataframe-to-word-document/40597684
	doc = Document('template_for_table.docx')
	# Change to landscape:
	# section = doc.sections[-1]
	# new_width, new_height = section.page_height, section.page_width
	# section.orientation = WD_ORIENT.LANDSCAPE
	# section.page_width = new_width
	# section.page_height = new_height

	date_now = str(datetime.now()).split('.')[0]
	
	heading = 'Pending Visits - ' + date_now
	print(heading)
	doc.add_heading(heading, 0)
	# add a table to the end and create a reference variable
	# extra row is so we can add the header row
	t = doc.add_table(df.shape[0]+1, df.shape[1])
	t.style = 'TableGrid'
	# add the header rows.
	for j in range(df.shape[-1]):
		cell = t.cell(0,j)
		cell.text = df.columns[j]
		run = cell.paragraphs[0].runs[0]
		run.font.bold = True

	# add the rest of the data frame
	for i in range(df.shape[0]):
		for j in range(df.shape[-1]):
			t.cell(i+1,j).text = str(df.values[i,j])

	# save the doc
	doc.save(filename)



# filename = askopenfilename()
filename = r'\\filestore.soton.ac.uk\users\clf1v16\mydesktop\small_example.csv'
data = pd.read_csv(filename)
# print(data.head())
df_columns = data.columns
new_df_columns = list(df_columns)
new_df_columns += ['Plans', 'Postcode', 'Check']

full_df_list = []
unique_names = []

for index, row in data.iterrows():
	row_list = list(row)
	row_list[5] = str(row_list[5])
	row_list[6] = str(row_list[6])
	# print(row_list)
	name = str(row["Patient"])
	full_care_plan = (str(row["Comments"]) + '\n' + str(row["Linked Care Plans"]))

	if name in unique_names:
		for p_record in full_df_list:
			if name == p_record[3]:
				p_record[5] += '\n' + str(full_care_plan)
				p_record[-3] += 1
				break
		continue
	else:
		unique_names.append(name)


	address = row["Patient Details"]
	# print(address)
	postcode = pf.string_to_postcode(address)

	cross_check_value = ''

	if 'insulin' in full_care_plan.lower():
		cross_check_value += 'INS'

	if 'clexane' in full_care_plan.lower():
		cross_check_value += 'CLE'

	row_list += [1, postcode, cross_check_value]

	full_df_list.append(copy(row_list))

df=pd.DataFrame(full_df_list,columns=new_df_columns)

word_filename = 'test_word_doc.docx'
print('Generating file: "' + word_filename + '"...')
df_to_word_table(word_filename, df)
print('Done.')
fname = os.path.basename(sys.argv[0])
cdir = sys.argv[0].split(fname)[0]
word_path = str(os.path.join(cdir, word_filename))
print(word_path)
os.system('START "" "' + word_path + '"')
# run(['open', word_filename], check=True)
# print(df)